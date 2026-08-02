"""Build the final reflective report DOCX from its Markdown source.

The Markdown remains the editable source of truth.  This script creates the
submission-friendly Word version with deterministic page, typography, list,
table, header, and footer settings.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(32, 55, 72)
GRAY = RGBColor(90, 98, 108)
LIGHT_FILL = "F4F6F9"
TABLE_HEADER_FILL = "F4F6F9"
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must total {TABLE_WIDTH_DXA}: {widths}")

    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_begin, instr, fld_sep, text, fld_end))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    # Named academic-report override: compact the narrative preset slightly so
    # the complete submission remains within the brief's 5-10 page target.
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("IncidentIQ | Reflective Report")
    set_run_font(hr, size=9, color=GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    add_page_field(fp)



INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|(?<!\*)\*[^*]+?\*(?!\*)|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str, *, default_size: float = 11) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=CODE_FONT, size=max(default_size - 0.5, 8.5))
        elif token.startswith("["):
            label = token[1 : token.index("]")].replace("`", "")
            url = token[token.index("(") + 1 : -1]
            display = label if label == url else f"{label} ({url})"
            run = paragraph.add_run(display)
            set_run_font(run, size=default_size, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=default_size)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(105)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("FINAL PROJECT REFLECTIVE REPORT"), size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("IncidentIQ"), size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    set_run_font(
        subtitle.add_run("Critical AI Use in Incident Response and Root-Cause Analysis"),
        size=15,
        color=DARK_BLUE,
    )

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course.paragraph_format.space_after = Pt(76)
    set_run_font(
        course.add_run("Computer Science - Critical Thinking, Problem Solving, and 21st-Century Skills"),
        size=10.5,
        color=GRAY,
        italic=True,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    set_run_font(meta.add_run("Development record: commits 1-11"), size=11, color=NAVY, bold=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta2.add_run("Evidence-first design | Human judgment remains in control"), size=10, color=GRAY)

    doc.add_page_break()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def table_widths(column_count: int) -> list[int]:
    if column_count == 2:
        return [2700, 6660]
    if column_count == 3:
        return [2160, 3600, 3600]
    base = TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(column_count))
    for row_index, source_row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            text = source_row[column_index] if column_index < len(source_row) else ""
            add_inline(paragraph, text, default_size=9.5)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
            if row_index == 0:
                shade_cell(cell, TABLE_HEADER_FILL)
        if row_index == 0:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: Iterable[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shd)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name=CODE_FONT, size=8.5, color=NAVY)


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    first_title_seen = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            index += 1
            add_code_block(doc, code_lines)
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            hashes, text = heading.groups()
            if len(hashes) == 1 and not first_title_seen:
                first_title_seen = True
                index += 1
                continue
            style = {1: "Heading 1", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[len(hashes)]
            paragraph = doc.add_paragraph(style=style)
            add_inline(paragraph, text, default_size={"Heading 1": 16, "Heading 2": 13, "Heading 3": 12}[style])
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet or numbered:
            paragraph = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_FILL)
            p_pr.append(shd)
            add_inline(paragraph, " ".join(quote_lines))
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = DARK_BLUE
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate == "---"
                or candidate.startswith(("#", "|", ">", "```"))
                or re.match(r"^[-*]\s+", candidate)
                or re.match(r"^\d+[.)]\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def build(source: Path, output: Path) -> None:
    markdown = source.read_text(encoding="utf-8")
    if any(marker in markdown for marker in ("«", "TODO", "TBD", "PLACEHOLDER")):
        raise ValueError("The report still contains a placeholder marker.")

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    render_markdown(doc, markdown)

    doc.core_properties.title = "IncidentIQ - Reflective Report"
    doc.core_properties.subject = "Critical AI use in incident response"
    doc.core_properties.keywords = "IncidentIQ, AI, incident response, critical thinking"
    doc.core_properties.comments = "Generated from docs/REFLECTIVE_REPORT.md"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
